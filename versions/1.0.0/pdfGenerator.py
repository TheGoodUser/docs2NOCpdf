
from docx import Document # type: ignore 
import docxedit # type: ignore


def generatePDF(FILLfileNo, FILLcName, FILLvehicleNo, FILLchasisNo, FILLengineNo, FILLDOC):
    document = Document('sample.docx')

    fileNo = document.tables[0].cell(1, 0)
    fileNo.text = ''  # Clear existing text
    fileNo.paragraphs[0].add_run(f'\n{FILLfileNo}\n').bold = True

    cName = document.tables[0].cell(1, 1)
    cName.text = ''  # Clear existing text
    cName.paragraphs[0].add_run(f'\n{FILLcName}\n').bold = True

    vehicleNo = document.tables[0].cell(1, 2)
    vehicleNo.text = ''  # Clear existing text
    vehicleNo.paragraphs[0].add_run(f'\n{FILLvehicleNo}\n').bold = True

    chasisNo = document.tables[0].cell(1, 3)
    chasisNo.text = ''  # Clear existing text
    chasisNo.paragraphs[0].add_run(f'\n{FILLchasisNo}\n').bold = True

    engineNo = document.tables[0].cell(1, 4)
    engineNo.text = ''  # Clear existing text
    engineNo.paragraphs[0].add_run(f'\n{FILLengineNo}\n').bold = True

    # Date Of Closing
    DOC = document.tables[0].cell(1, 5)
    DOC.text = ''  # Clear existing text
    DOC.paragraphs[0].add_run(f'\n{FILLDOC}\n').bold = True

    
    document.save(fr'./outputs/{FILLfileNo}.docx')
    print(f"PDF generated for {FILLcName} with MBFL No. {FILLfileNo}")









