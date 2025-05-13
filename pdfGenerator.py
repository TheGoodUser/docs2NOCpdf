
from docx import Document # type: ignore 
import docxedit # type: ignore
from datetime import datetime as dt
from tkinter import messagebox as msgbx
# print()


def generatePDF(FILLfileNo, FILLcName, FILLvehicleNo, FILLchasisNo, FILLengineNo, FILLDOC):
    document = Document('sample.docx')
    todaysDate = dt.now().strftime(r'%d/%m/20%y')
    # ============================ Date ======================================
    date = document.paragraphs[3]
    date.clear()
    date_run = date.add_run(f"TO,\t\t\t\t\t\t\t\t\t\t\t\tDATE: {todaysDate}")
    date_run.bold = True
    
    

    # ============== Subject Section =========================================
    subject = document.paragraphs[7]
    subject.clear()
    # SUB: NOC FOR ACCOUNT NO-MBFL-1231, HYPOTHECATION FOR VEHICLE NUMBER -CG-22-U-6546
    run = subject.add_run(f"SUB: NOC FOR ACCOUNT NO.- {FILLfileNo}, HYPOTHECATION FOR VEHICLE NUMBER- {FILLvehicleNo}")
    run.bold = True



    # ============== Table section =========================================
    fileNo = document.tables[0].cell(1, 0)
    fileNo.text = ''  # Clear existing text
    fileNo.paragraphs[0].add_run(f'\n{FILLfileNo}\n').bold = True

    cName = document.tables[0].cell(1, 1)
    cName.text = ''  # Clear existing text
    cName.paragraphs[0].add_run(f'\n{FILLcName}\n').bold = True

    vehicleNo = document.tables[0].cell(1, 2)
    vehicleNo.text = ''  # Clear existing text
    vehicleNo.paragraphs[0].add_run(f'\n{FILLvehicleNo}\n').bold = True

    chasisNo = document.tables[0].cell(1, 3)
    chasisNo.text = ''  # Clear existing text
    chasisNo.paragraphs[0].add_run(f'\n{FILLchasisNo}\n').bold = True

    engineNo = document.tables[0].cell(1, 4)
    engineNo.text = ''  # Clear existing text
    engineNo.paragraphs[0].add_run(f'\n{FILLengineNo}\n').bold = True

    # Date Of Closing
    DOC = document.tables[0].cell(1, 5)
    DOC.text = ''  # Clear existing text
    DOC.paragraphs[0].add_run(f'\n{FILLDOC}\n').bold = True

    
    try:
        document.save(fr'./outputs/{FILLfileNo}.docx')
    except:
        msgbx.showinfo("Excel2PDF", f"Close the Opened NOC file !!\nNamed: {FILLcName}\nFile No.: {FILLfileNo}\nDated: {todaysDate}")
        return 0

    print(f"PDF generated for {FILLcName} with MBFL No. {FILLfileNo}")









